from io import BytesIO

from docx import Document
from sqlalchemy.pool import StaticPool
from sqlmodel import Session, SQLModel, create_engine, select

from app.business_rules import get_business_rules
from app.db import get_session
from app.main import app
from app.models import BusinessRules, Item, ItemMaster, Supplier

try:
    from fastapi.testclient import TestClient
except ImportError:  # pragma: no cover
    TestClient = None


def _fresh_session() -> Session:
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    SQLModel.metadata.create_all(engine)
    return Session(engine)


def _make_client():
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    SQLModel.metadata.create_all(engine)

    def override_get_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_get_session
    return TestClient(app), engine


def test_get_business_rules_creates_default_singleton_once():
    with _fresh_session() as session:
        rules1 = get_business_rules(session)
        assert rules1.id == 1
        # Defaults match the values that used to be hardcoded constants.
        assert rules1.security_deposit_percent == 5.0
        assert rules1.security_deposit_waived_below == 0.0
        assert rules1.stamp_duty_percent == 0.25

        rules2 = get_business_rules(session)
        assert rules2.id == rules1.id

        assert len(session.exec(select(BusinessRules)).all()) == 1


def test_settings_page_updates_persist_and_affect_next_contract_award():
    client, engine = _make_client()
    try:
        # Set up a tender with one awarded item via the real HTTP routes.
        resp = client.post("/tenders", data={"inquiry_no": "Rules Test", "tax_percent": "10"}, follow_redirects=False)
        tender_id = int(resp.headers["location"].rsplit("/", 1)[-1])

        with Session(engine) as session:
            im = ItemMaster(part_no="X-1", description="Widget", default_unit="Nos")
            session.add(im)
            session.commit()
            session.refresh(im)
            item_master_id = im.id

        client.post(f"/tenders/{tender_id}/items", data={"item_master_id": str(item_master_id), "qty": "5"})
        with Session(engine) as session:
            item_id = session.exec(select(Item).where(Item.tender_id == tender_id)).one().id

        supplier_id = client.post("/suppliers/quick-create", data={"name": "Acme"}).json()["id"]
        client.post(
            f"/tenders/{tender_id}/quote-entry",
            data={"supplier_id": str(supplier_id), f"rate__{item_id}": "1000"},
        )
        client.post(f"/tenders/{tender_id}/generate-proposal")
        client.post(f"/tenders/{tender_id}/approve-proposal")

        with Session(engine) as session:
            supplier = session.get(Supplier, supplier_id)
            supplier.address = "Test Address"
            session.add(supplier)
            session.commit()

        # Default rules: 5% deposit always applies -> 5000*5% = 250.00
        resp = client.post(f"/tenders/{tender_id}/proposal/contract/{supplier_id}", data={"contract_no": "C-1"})
        assert resp.status_code == 200
        full_text = "\n".join(p.text for p in Document(BytesIO(resp.content)).paragraphs)
        assert "250.00" in full_text

        # Now set a threshold above this contract's value via the real
        # settings form - the deposit should disappear on the next doc.
        resp = client.post(
            "/settings/business-rules",
            data={
                "security_deposit_percent": "5",
                "security_deposit_waived_below": "999999",
                "stamp_duty_percent": "0.25",
            },
            follow_redirects=False,
        )
        assert resp.status_code == 303

        resp = client.post(f"/tenders/{tender_id}/proposal/contract/{supplier_id}", data={"contract_no": "C-2"})
        assert resp.status_code == 200
        full_text = "\n".join(p.text for p in Document(BytesIO(resp.content)).paragraphs)
        assert "250.00" not in full_text
        assert "0.00" in full_text
    finally:
        app.dependency_overrides.clear()
