import sys
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy.pool import StaticPool
from sqlmodel import Session, SQLModel, create_engine

from app import template_manager
from app.db import get_session
from app.main import app
from app.paths import resource_path

try:
    from fastapi.testclient import TestClient
except ImportError:  # pragma: no cover
    TestClient = None

REAL_CA_BYTES = resource_path("docx_templates", "ca_template.docx").read_bytes()
REAL_PP_BYTES = resource_path("docx_templates", "pp_template.docx").read_bytes()

# CA.doc is the user-supplied real reference document (see CLAUDE.md "Data
# sensitivity") - gitignored, local-only. Tests that need real .doc content
# skip gracefully when it isn't present rather than failing on a machine/CI
# that doesn't have it.
REAL_DOC_PATH = Path(__file__).resolve().parent.parent / "CA.doc"
requires_real_doc_file = pytest.mark.skipif(not REAL_DOC_PATH.exists(), reason="CA.doc not present (local-only sample file)")


# Isolation from the real custom_docx_templates_dir() is provided
# suite-wide by tests/conftest.py's autouse fixture.


def _make_client():
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    SQLModel.metadata.create_all(engine)

    def override_get_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_get_session
    return TestClient(app), engine


def test_list_templates_reports_default_until_a_custom_upload_exists():
    rows = template_manager.list_templates()
    by_name = {r["name"]: r for r in rows}
    assert set(by_name) == {"ca_template.docx", "pp_template.docx"}
    assert by_name["ca_template.docx"]["is_custom"] is False

    template_manager.save_custom_template("ca_template.docx", REAL_CA_BYTES)
    rows = template_manager.list_templates()
    by_name = {r["name"]: r for r in rows}
    assert by_name["ca_template.docx"]["is_custom"] is True
    assert by_name["pp_template.docx"]["is_custom"] is False


def test_save_then_restore_default_round_trips():
    template_manager.save_custom_template("ca_template.docx", REAL_CA_BYTES)
    assert template_manager.read_active_template("ca_template.docx") == REAL_CA_BYTES

    template_manager.restore_default_template("ca_template.docx")
    rows = {r["name"]: r for r in template_manager.list_templates()}
    assert rows["ca_template.docx"]["is_custom"] is False
    # Restoring twice (nothing to delete) must not raise.
    template_manager.restore_default_template("ca_template.docx")


def test_validate_template_accepts_real_templates_and_rejects_garbage():
    template_manager.validate_template("ca_template.docx", REAL_CA_BYTES)  # must not raise
    template_manager.validate_template("pp_template.docx", REAL_PP_BYTES)

    with pytest.raises(ValueError, match="Contract Award"):
        template_manager.validate_template("ca_template.docx", b"not a docx")


def test_validate_template_rejects_broken_jinja_tag():
    doc = Document(BytesIO(REAL_CA_BYTES))
    doc.paragraphs[0].insert_paragraph_before("{% if unclosed %}broken")
    buffer = BytesIO()
    doc.save(buffer)

    with pytest.raises(ValueError):
        template_manager.validate_template("ca_template.docx", buffer.getvalue())


def test_unknown_template_name_is_rejected_everywhere():
    with pytest.raises(ValueError):
        template_manager.save_custom_template("../../evil.docx", b"x")
    with pytest.raises(ValueError):
        template_manager.read_active_template("not_a_real_template.docx")
    with pytest.raises(ValueError):
        template_manager.validate_template("not_a_real_template.docx", b"x")


def test_upload_route_rejects_invalid_file_and_accepts_valid_one():
    client, engine = _make_client()
    try:
        resp = client.post(
            "/settings/templates/ca_template.docx/upload",
            files={"file": ("bad.docx", b"not a docx", "application/octet-stream")},
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert "error=" in resp.headers["location"]

        resp = client.get(resp.headers["location"])
        assert "Could not use this file" in resp.text

        resp = client.post(
            "/settings/templates/ca_template.docx/upload",
            files={"file": ("ca_template.docx", REAL_CA_BYTES, "application/octet-stream")},
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert resp.headers["location"] == "/settings/templates?saved=1"

        resp = client.get("/settings/templates/ca_template.docx/download")
        assert resp.status_code == 200
        assert resp.content == REAL_CA_BYTES
    finally:
        app.dependency_overrides.clear()


@requires_real_doc_file
def test_convert_doc_to_docx_produces_a_valid_docx():
    content = REAL_DOC_PATH.read_bytes()
    result = template_manager.convert_doc_to_docx(content)
    assert result[:2] == b"PK"  # zip signature - a real .docx, not raw .doc bytes
    # Must be openable as a real docx, not just superficially zip-shaped.
    Document(BytesIO(result))


def test_convert_doc_to_docx_gives_a_friendly_error_when_word_unavailable(monkeypatch):
    monkeypatch.setitem(sys.modules, "win32com", None)
    monkeypatch.setitem(sys.modules, "win32com.client", None)
    monkeypatch.setitem(sys.modules, "pythoncom", None)

    with pytest.raises(ValueError, match="Microsoft Word"):
        template_manager.convert_doc_to_docx(b"pretend .doc bytes")


@requires_real_doc_file
def test_upload_route_converts_a_real_doc_upload():
    client, engine = _make_client()
    try:
        resp = client.post(
            "/settings/templates/ca_template.docx/upload",
            files={"file": ("CA.doc", REAL_DOC_PATH.read_bytes(), "application/msword")},
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert resp.headers["location"] == "/settings/templates?saved=1"

        # What actually got stored/is now active is a real, valid .docx -
        # not the raw uploaded .doc bytes.
        resp = client.get("/settings/templates/ca_template.docx/download")
        assert resp.status_code == 200
        assert resp.content[:2] == b"PK"
        Document(BytesIO(resp.content))
    finally:
        app.dependency_overrides.clear()


def test_upload_route_rejects_unsupported_extension():
    client, engine = _make_client()
    try:
        resp = client.post(
            "/settings/templates/ca_template.docx/upload",
            files={"file": ("template.txt", b"plain text", "text/plain")},
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert "error=" in resp.headers["location"]
        resp = client.get(resp.headers["location"])
        assert ".doc" in resp.text
    finally:
        app.dependency_overrides.clear()
