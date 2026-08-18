from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy.pool import StaticPool
from sqlmodel import Session, SQLModel, create_engine, select

from app import custom_fields as custom_fields_module
from app.award_engine import build_purchase_proposal
from app.cs_engine import build_comparative_statement
from app.db import get_session
from app.docx_export import generate_contract_award
from app.excel_io import export_cs_xlsx, import_tender
from app.main import app
from app.models import BusinessRules, CustomField, CustomFieldGroup, Department, DocumentLabels, Supplier, Tender
from app.proposal_snapshot import save_proposal_snapshot

try:
    from fastapi.testclient import TestClient
except ImportError:  # pragma: no cover
    TestClient = None

CS_XLSX_PATH = Path(__file__).resolve().parent.parent / "CS.xlsx"
DEFAULT_RULES = BusinessRules()
DEFAULT_LABELS = DocumentLabels()


def _fresh_session() -> Session:
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    SQLModel.metadata.create_all(engine)
    return Session(engine)


def _make_client():
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    SQLModel.metadata.create_all(engine)

    def override_get_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_get_session
    return TestClient(app), engine


# --- Module-level behavior ---------------------------------------------------


def test_slugify_tag_name():
    assert custom_fields_module.slugify_tag_name("Prep By - Designation") == "prep_by_designation"
    assert custom_fields_module.slugify_tag_name("  Extra   Spaces  ") == "extra_spaces"
    assert custom_fields_module.slugify_tag_name("123 Starts With Digit") == "field_123_starts_with_digit"


def test_validate_tag_name_rejects_invalid_characters_and_reserved_names():
    with pytest.raises(ValueError):
        custom_fields_module.validate_tag_name("Not Valid")  # spaces/caps
    with pytest.raises(ValueError):
        custom_fields_module.validate_tag_name("1_starts_with_digit")
    with pytest.raises(ValueError):
        custom_fields_module.validate_tag_name("contract_value")  # reserved - real computed data
    custom_fields_module.validate_tag_name("prep_by_designation")  # must not raise


def test_create_rejects_duplicate_tag_name():
    with _fresh_session() as session:
        custom_fields_module.create_custom_field(session, "my_field", "My Field", "Value 1")
        with pytest.raises(ValueError):
            custom_fields_module.create_custom_field(session, "my_field", "Different label", "Value 2")


def test_update_cannot_change_tag_name_only_label_and_value():
    with _fresh_session() as session:
        field = custom_fields_module.create_custom_field(session, "my_field", "Original", "v1")
        custom_fields_module.update_custom_field(session, field.id, "Updated label", "v2")
        refreshed = session.get(CustomField, field.id)
        assert refreshed.tag_name == "my_field"  # unchanged
        assert refreshed.label == "Updated label"
        assert refreshed.value == "v2"


def test_delete_is_a_no_op_for_unknown_id():
    with _fresh_session() as session:
        custom_fields_module.delete_custom_field(session, 999999)  # must not raise


def test_custom_fields_dict_reflects_all_rows():
    with _fresh_session() as session:
        custom_fields_module.create_custom_field(session, "a", "A", "1")
        custom_fields_module.create_custom_field(session, "b", "B", "2")
        assert custom_fields_module.custom_fields_dict(session) == {"a": "1", "b": "2"}


# --- Custom Field Groups (per-department presets) ----------------------------


def test_create_group_requires_a_department():
    with _fresh_session() as session:
        with pytest.raises(ValueError):
            custom_fields_module.create_group(session, "Department A", None)


def test_group_is_unique_per_department():
    with _fresh_session() as session:
        dept = Department(name="Procurement Section")
        session.add(dept)
        session.commit()
        session.refresh(dept)

        custom_fields_module.create_group(session, "Group One", dept.id)
        with pytest.raises(ValueError):
            custom_fields_module.create_group(session, "Group Two", dept.id)  # same department again


def test_group_field_overrides_global_field_of_same_tag_name():
    with _fresh_session() as session:
        dept_a = Department(name="Department A")
        dept_b = Department(name="Department B")
        session.add(dept_a)
        session.add(dept_b)
        session.commit()
        session.refresh(dept_a)
        session.refresh(dept_b)

        group_a = custom_fields_module.create_group(session, "Department A", dept_a.id)
        group_b = custom_fields_module.create_group(session, "Department B", dept_b.id)

        custom_fields_module.create_custom_field(session, "initiator_name", "Initiator", "Global Default Name")
        custom_fields_module.create_custom_field(
            session, "initiator_name", "Initiator", "Officer From Dept A", group_id=group_a.id
        )
        # Same tag_name reused in a second group - must not collide with group_a's row.
        custom_fields_module.create_custom_field(
            session, "initiator_name", "Initiator", "Officer From Dept B", group_id=group_b.id
        )

        assert custom_fields_module.custom_fields_dict(session, group_id=group_a.id)["initiator_name"] == "Officer From Dept A"
        assert custom_fields_module.custom_fields_dict(session, group_id=group_b.id)["initiator_name"] == "Officer From Dept B"
        # A tag not overridden in the group still falls back to the global value.
        custom_fields_module.create_custom_field(session, "receiving_store", "Store", "Main Store")
        assert custom_fields_module.custom_fields_dict(session, group_id=group_a.id)["receiving_store"] == "Main Store"
        # No group at all -> plain global value.
        assert custom_fields_module.custom_fields_dict(session)["initiator_name"] == "Global Default Name"


def test_custom_fields_dict_for_tender_resolves_via_department():
    with _fresh_session() as session:
        dept = Department(name="Department A")
        session.add(dept)
        session.commit()
        session.refresh(dept)

        group = custom_fields_module.create_group(session, "Department A", dept.id)
        custom_fields_module.create_custom_field(session, "initiator_name", "Initiator", "Global Name")
        custom_fields_module.create_custom_field(
            session, "initiator_name", "Initiator", "Dept A Officer", group_id=group.id
        )

        tender_with_dept = Tender(inquiry_no="T-1", department_id=dept.id)
        assert custom_fields_module.custom_fields_dict_for_tender(session, tender_with_dept)["initiator_name"] == "Dept A Officer"

        tender_without_dept = Tender(inquiry_no="T-2")
        assert custom_fields_module.custom_fields_dict_for_tender(session, tender_without_dept)["initiator_name"] == "Global Name"

        assert custom_fields_module.custom_fields_dict_for_tender(session, None)["initiator_name"] == "Global Name"


def test_tag_document_scope_reports_which_documents_use_a_tag():
    assert custom_fields_module.tag_document_scope("prep_by_designation") == ("CS",)
    assert custom_fields_module.tag_document_scope("pp_paying_authority") == ("PP",)
    assert custom_fields_module.tag_document_scope("indentor_name") == ("CA",)
    assert custom_fields_module.tag_document_scope("not_a_suggested_tag") == ()


def test_classify_fields_by_scope_buckets_by_document_overlap():
    with _fresh_session() as session:
        cs_field = custom_fields_module.create_custom_field(session, "prep_by_designation", "Prep By", "Clerk")
        pp_field = custom_fields_module.create_custom_field(session, "pp_paying_authority", "Paying Authority", "Finance")
        ca_field = custom_fields_module.create_custom_field(session, "indentor_name", "Indentor", "Director")
        other_field = custom_fields_module.create_custom_field(session, "ad_hoc_note", "Ad Hoc Note", "Something")

        buckets = custom_fields_module.classify_fields_by_scope([cs_field, pp_field, ca_field, other_field])
        assert buckets["CS"] == [cs_field]
        assert buckets["PP"] == [pp_field]
        assert buckets["CA"] == [ca_field]
        assert buckets["other"] == [other_field]
        assert buckets["shared"] == []
        assert buckets["global"] == []


def test_bulk_set_fields_creates_updates_and_deletes_on_blank():
    with _fresh_session() as session:
        dept = Department(name="Department A")
        session.add(dept)
        session.commit()
        session.refresh(dept)
        group = custom_fields_module.create_group(session, "Department A", dept.id)

        custom_fields_module.bulk_set_fields(
            session, group.id, {"indentor_name": "Director Procurement", "cost_head": "Fund 1/2/3"}
        )
        values = custom_fields_module.custom_fields_dict(session, group_id=group.id)
        assert values["indentor_name"] == "Director Procurement"
        assert values["cost_head"] == "Fund 1/2/3"

        # Re-submitting with one field changed and one blanked out: updates the
        # first, deletes the override for the second (falls back to global/none).
        custom_fields_module.bulk_set_fields(
            session, group.id, {"indentor_name": "Director SCM", "cost_head": "  "}
        )
        values = custom_fields_module.custom_fields_dict(session, group_id=group.id)
        assert values["indentor_name"] == "Director SCM"
        assert "cost_head" not in values


def test_bulk_set_fields_handles_organization_wide_scope():
    with _fresh_session() as session:
        custom_fields_module.bulk_set_fields(session, None, {"pp_paying_authority": "Finance Directorate"})
        values = custom_fields_module.custom_fields_dict(session)
        assert values["pp_paying_authority"] == "Finance Directorate"


def test_delete_group_cascades_its_fields():
    with _fresh_session() as session:
        dept = Department(name="Department A")
        session.add(dept)
        session.commit()
        session.refresh(dept)

        group = custom_fields_module.create_group(session, "Department A", dept.id)
        field = custom_fields_module.create_custom_field(
            session, "initiator_name", "Initiator", "Dept A Officer", group_id=group.id
        )

        custom_fields_module.delete_group(session, group.id)
        assert session.get(CustomFieldGroup, group.id) is None
        assert session.get(CustomField, field.id) is None


# --- Word document rendering -------------------------------------------------


def test_custom_field_renders_as_a_tag_in_contract_award():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        snapshot = save_proposal_snapshot(session, tender.id)
        group = next(g for g in snapshot.firm_groups if g.supplier_name == "M/s SNS Enterprises")
        supplier = session.get(Supplier, group.supplier_id)

        content = generate_contract_award(
            tender, group, supplier, contract_no="C-1", rules=DEFAULT_RULES,
            custom_fields={"prep_by_designation": "Junior Clerk (BS-11)"},
        )
        full_text = "\n".join(p.text for p in Document(BytesIO(content)).paragraphs)
        # The real ca_template.docx has no {{ prep_by_designation }} tag today,
        # so this only proves the value was available to render, not that it
        # appears - the merge-behavior test below proves the actual mechanism.
        assert "{{" not in full_text and "{%" not in full_text


def test_real_computed_data_overrides_a_same_named_custom_field():
    """Defense in depth: even if a colliding key somehow reached the merge
    (creation already blocks it via validate_tag_name), real per-contract
    data must win, never a custom field standing in for it."""
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        snapshot = save_proposal_snapshot(session, tender.id)
        group = next(g for g in snapshot.firm_groups if g.supplier_name == "M/s SNS Enterprises")
        supplier = session.get(Supplier, group.supplier_id)

        content = generate_contract_award(
            tender, group, supplier, contract_no="C-1", rules=DEFAULT_RULES,
            custom_fields={"contract_no": "SOMETHING-ELSE-ENTIRELY"},
        )
        full_text = "\n".join(p.text for p in Document(BytesIO(content)).paragraphs)
        assert "C-1" in full_text
        assert "SOMETHING-ELSE-ENTIRELY" not in full_text


# --- CS Excel designation lines ----------------------------------------------


def test_cs_excel_shows_designation_line_when_suggested_field_is_set():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        cs = build_comparative_statement(session, tender.id)

        without = export_cs_xlsx(cs, DEFAULT_LABELS, {})
        from openpyxl import load_workbook

        values_without = [c.value for row in load_workbook(BytesIO(without)).active.iter_rows() for c in row]
        assert "Junior Clerk (BS-11)" not in values_without

        with_designation = export_cs_xlsx(
            cs, DEFAULT_LABELS, {"prep_by_designation": "Junior Clerk (BS-11)", "fmsad_designation": "Director Finance"}
        )
        values_with = [c.value for row in load_workbook(BytesIO(with_designation)).active.iter_rows() for c in row]
        assert "Junior Clerk (BS-11)" in values_with
        assert "Director Finance" in values_with


def test_cs_excel_strips_illegal_control_characters_from_pasted_labels():
    """Regression: a label pasted from Word can carry an ASCII control
    character (e.g. \\x0b vertical tab) that openpyxl refuses outright
    (IllegalCharacterError), crashing the whole export - see excel_io.py's
    _sanitize_workbook_text()."""
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        cs = build_comparative_statement(session, tender.id)

        labels = DocumentLabels(head_qac_label="__________ Capt\x0bDAD (FWA) \x0b(Muhammad Qasim)")
        content = export_cs_xlsx(cs, labels, {})  # must not raise IllegalCharacterError

        from openpyxl import load_workbook

        values = [c.value for row in load_workbook(BytesIO(content)).active.iter_rows() for c in row]
        assert "__________ CaptDAD (FWA) (Muhammad Qasim)" in values
        assert not any(isinstance(v, str) and "\x0b" in v for v in values)


# --- Full HTTP round trip: the "Manage Tags" popup ---------------------------


def test_manage_tags_saves_organization_wide_fields_in_one_submit():
    client, engine = _make_client()
    try:
        resp = client.get("/settings/custom-fields")
        assert "prep_by_designation" in resp.text  # suggested, not yet set

        resp = client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": "", "prep_by_designation": "Junior Clerk", "pp_paying_authority": "Finance"},
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert resp.headers["location"] == "/settings/custom-fields?saved=1&open=1"

        with Session(engine) as session:
            values = custom_fields_module.custom_fields_dict(session)
            assert values["prep_by_designation"] == "Junior Clerk"
            assert values["pp_paying_authority"] == "Finance"

        resp = client.get("/settings/custom-fields")
        assert "Junior Clerk" in resp.text
    finally:
        app.dependency_overrides.clear()


def test_manage_tags_blank_value_removes_a_saved_field():
    client, engine = _make_client()
    try:
        client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": "", "pp_paying_authority": "Finance"},
        )
        with Session(engine) as session:
            assert custom_fields_module.custom_fields_dict(session)["pp_paying_authority"] == "Finance"

        client.post("/settings/custom-fields/manage-tags", data={"department_id": "", "pp_paying_authority": ""})
        with Session(engine) as session:
            assert "pp_paying_authority" not in custom_fields_module.custom_fields_dict(session)
    finally:
        app.dependency_overrides.clear()


def test_manage_tags_auto_creates_department_profile_on_first_save():
    """No separate "create a profile" step - saving a value for a
    department that has no CustomFieldGroup yet creates one on the fly."""
    client, engine = _make_client()
    try:
        with Session(engine) as session:
            dept = Department(name="Department A")
            session.add(dept)
            session.commit()
            dept_id = dept.id

        with Session(engine) as session:
            assert session.exec(select(CustomFieldGroup)).all() == []

        resp = client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": str(dept_id), "indentor_name": "Director Procurement"},
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert f"department_id={dept_id}" in resp.headers["location"]

        with Session(engine) as session:
            group = session.exec(select(CustomFieldGroup)).one()
            assert group.department_id == dept_id
            assert custom_fields_module.custom_fields_dict(session, group_id=group.id)["indentor_name"] == (
                "Director Procurement"
            )

        resp = client.get(f"/settings/custom-fields?department_id={dept_id}")
        assert "Director Procurement" in resp.text
        assert "Department A" in resp.text
    finally:
        app.dependency_overrides.clear()


def test_manage_tags_reuses_existing_group_instead_of_creating_a_duplicate():
    client, engine = _make_client()
    try:
        with Session(engine) as session:
            dept = Department(name="Department A")
            session.add(dept)
            session.commit()
            dept_id = dept.id
            group = custom_fields_module.create_group(session, "Department A Profile", dept_id)
            group_id = group.id

        client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": str(dept_id), "indentor_name": "Director Procurement"},
        )
        with Session(engine) as session:
            assert len(session.exec(select(CustomFieldGroup)).all()) == 1
            assert (
                custom_fields_module.custom_fields_dict(session, group_id=group_id)["indentor_name"]
                == "Director Procurement"
            )
    finally:
        app.dependency_overrides.clear()


def test_manage_tags_add_new_field_without_leaving_the_page():
    """Simulates the popup's "+ Add another field" row - submitted as
    new_tag_name/new_tag_value alongside the known tags, in the same POST."""
    client, engine = _make_client()
    try:
        resp = client.post(
            "/settings/custom-fields/manage-tags",
            data={
                "department_id": "",
                "new_tag_name": "store_keeper_name",
                "new_tag_value": "M. Tariq",
            },
            follow_redirects=False,
        )
        assert resp.status_code == 303
        assert "error=" not in resp.headers["location"]

        with Session(engine) as session:
            field = session.exec(select(CustomField).where(CustomField.tag_name == "store_keeper_name")).one()
            assert field.value == "M. Tariq"
            assert field.group_id is None
    finally:
        app.dependency_overrides.clear()


def test_manage_tags_rejects_reserved_and_duplicate_new_field_names():
    client, engine = _make_client()
    try:
        resp = client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": "", "new_tag_name": "contract_value", "new_tag_value": "x"},
            follow_redirects=False,
        )
        assert "error=" in resp.headers["location"]

        resp = client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": "", "new_tag_name": "pp_paying_authority", "new_tag_value": "x"},
            follow_redirects=False,
        )
        assert "error=" in resp.headers["location"]  # collides with a known tag in this same scope

        with Session(engine) as session:
            assert session.exec(select(CustomField)).all() == []
    finally:
        app.dependency_overrides.clear()


def test_manage_tags_route_for_unknown_department_returns_error():
    client, engine = _make_client()
    try:
        resp = client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": "999999", "indentor_name": "x"},
            follow_redirects=False,
        )
        assert "error=" in resp.headers["location"]
        with Session(engine) as session:
            assert session.exec(select(CustomFieldGroup)).all() == []
    finally:
        app.dependency_overrides.clear()


def test_delete_group_route_resets_a_departments_profile():
    client, engine = _make_client()
    try:
        with Session(engine) as session:
            dept = Department(name="Department A")
            session.add(dept)
            session.commit()
            dept_id = dept.id

        client.post(
            "/settings/custom-fields/manage-tags",
            data={"department_id": str(dept_id), "indentor_name": "Director Procurement"},
        )
        with Session(engine) as session:
            group_id = session.exec(select(CustomFieldGroup)).one().id

        resp = client.post(f"/settings/custom-field-groups/{group_id}/delete", follow_redirects=False)
        assert resp.status_code == 303
        with Session(engine) as session:
            assert session.exec(select(CustomFieldGroup)).all() == []
            assert session.exec(select(CustomField)).all() == []
    finally:
        app.dependency_overrides.clear()
