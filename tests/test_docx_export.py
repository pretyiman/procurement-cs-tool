import datetime
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from sqlalchemy.pool import StaticPool
from sqlmodel import Session, SQLModel, create_engine

from app.award_engine import build_purchase_proposal
from app.docx_export import generate_contract_award, generate_purchase_proposal_doc
from app.excel_io import import_tender
from app.models import BusinessRules, Supplier
from app.proposal_snapshot import save_proposal_snapshot

CS_XLSX_PATH = Path(__file__).resolve().parent.parent / "CS.xlsx"
DEFAULT_RULES = BusinessRules()  # 5% deposit (never waived), 0.25% stamp duty - matches the old hardcoded constants


def _fresh_session() -> Session:
    engine = create_engine(
        "sqlite://", connect_args={"check_same_thread": False}, poolclass=StaticPool
    )
    SQLModel.metadata.create_all(engine)
    return Session(engine)


def _full_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def _snapshot_and_group(session: Session, tender, supplier_name: str):
    """Generate-and-freeze a proposal snapshot (the real path every
    Contract Award/Purchase Proposal document renders from now), and pull
    out one firm's frozen group by name."""
    snapshot = save_proposal_snapshot(session, tender.id)
    group = next(g for g in snapshot.firm_groups if g.supplier_name == supplier_name)
    return snapshot, group


# --- Contract Award (CA) ----------------------------------------------------


def test_contract_award_item_schedule_matches_proposal_exactly():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        # Live proposal (pre-snapshot) is still the ground truth to compare
        # the frozen snapshot's rendered numbers against.
        live_proposal = build_purchase_proposal(session, tender.id)
        live_group = next(g for g in live_proposal.firm_groups if g.supplier_name == "M/s SNS Enterprises")

        _snapshot, group = _snapshot_and_group(session, tender, "M/s SNS Enterprises")
        supplier = session.get(Supplier, group.supplier_id)
        supplier.address = "Test Address"
        session.add(supplier)
        session.commit()

        content = generate_contract_award(
            tender, group, supplier, contract_no="TEST-001", rules=DEFAULT_RULES,
            contract_date=datetime.date(2026, 8, 12), agreement_date=datetime.date(2026, 8, 12),
        )
        doc = Document(BytesIO(content))

        # tables[0] is the static "WARNING" notice, tables[1] is the item schedule.
        table = doc.tables[1]
        assert len(table.rows) == 1 + len(group.items) + 3  # header + items + 3 totals rows

        rendered_rows = [tuple(c.text for c in row.cells) for row in table.rows[1 : 1 + len(group.items)]]
        expected_rows = [
            (
                str(ai.item.ser),
                ai.item.item_master.part_no,
                ai.item.item_master.description,
                ai.item.item_master.default_unit,
                str(ai.item.qty),
                f"{ai.awarded_rate:,.2f}",
                f"{ai.total_value:,.2f}",
            )
            for ai in live_group.items
        ]
        assert rendered_rows == expected_rows

        full_text = _full_text(doc)
        assert f"{group.store_value:,.2f}" in full_text
        assert f"{group.contract_value:,.2f}" in full_text
        assert "TEST-001" in full_text
        assert "{{" not in full_text and "{%" not in full_text


def test_contract_award_includes_opening_date():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        tender.opening_date = datetime.date(2026, 7, 15)
        session.add(tender)
        session.commit()

        _snapshot, group = _snapshot_and_group(session, tender, "M/s SNS Enterprises")
        supplier = session.get(Supplier, group.supplier_id)

        content = generate_contract_award(tender, group, supplier, contract_no="C-OPEN", rules=DEFAULT_RULES)
        full_text = _full_text(Document(BytesIO(content)))
        assert "15 Jul 2026" in full_text
        assert "{{" not in full_text and "{%" not in full_text


def test_contract_award_opening_date_placeholder_when_unset():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        _snapshot, group = _snapshot_and_group(session, tender, "M/s SNS Enterprises")
        supplier = session.get(Supplier, group.supplier_id)

        content = generate_contract_award(tender, group, supplier, contract_no="C-2", rules=DEFAULT_RULES)
        full_text = _full_text(Document(BytesIO(content)))
        assert "Tender Opening Date ___" in full_text


def test_contract_award_amount_in_words_and_computed_fees():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        _snapshot, group = _snapshot_and_group(session, tender, "M/s Awan Tech")
        supplier = session.get(Supplier, group.supplier_id)

        content = generate_contract_award(tender, group, supplier, contract_no="C-1", rules=DEFAULT_RULES)
        full_text = _full_text(Document(BytesIO(content)))

        # Matches the real sample CA.doc's amount-in-words for this exact firm/value.
        assert "Pak Rupees Two Hundred Forty Nine Thousand One Hundred Thirty Eight and Paisa Twelve Only" in full_text
        # Both security deposit and stamp duty are on store value (excl.
        # tax), not contract value (incl. tax) - see docx_export.py.
        assert f"{group.store_value * 0.05:,.2f}" in full_text
        assert f"{group.store_value * 0.0025:,.2f}" in full_text


def test_security_deposit_waived_below_configured_threshold():
    """The exact scenario that motivated BusinessRules: a contract below a
    configured threshold gets no security deposit at all, while one at or
    above it still gets the normal percentage."""
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        _snapshot, group = _snapshot_and_group(session, tender, "M/s Awan Tech")
        supplier = session.get(Supplier, group.supplier_id)

        # Threshold above this firm's contract value -> deposit waived (0.00).
        waived_rules = BusinessRules(security_deposit_percent=5.0, security_deposit_waived_below=10_000_000, stamp_duty_percent=0.25)
        content = generate_contract_award(tender, group, supplier, contract_no="C-WAIVED", rules=waived_rules)
        full_text = _full_text(Document(BytesIO(content)))
        assert "0.00" in full_text
        assert f"{group.store_value * 0.05:,.2f}" not in full_text

        # Threshold below this firm's contract value -> normal 5% still applies.
        applies_rules = BusinessRules(security_deposit_percent=5.0, security_deposit_waived_below=1, stamp_duty_percent=0.25)
        content = generate_contract_award(tender, group, supplier, contract_no="C-APPLIES", rules=applies_rules)
        full_text = _full_text(Document(BytesIO(content)))
        assert f"{group.store_value * 0.05:,.2f}" in full_text


def test_stamp_duty_and_security_deposit_are_both_based_on_store_value():
    """Procurement rule: stamp duty and the security deposit (bank
    guarantee) are both calculated on store value (quoted price excl.
    GST/PST), never on contract value (incl. tax) - even though store_value
    != contract_value whenever tax_percent > 0, so this fixture (18% GST)
    would silently pass a contract-value-based calculation too if it
    matched by coincidence. Assert the store-value figure is present AND
    the contract-value figure is absent, so a regression can't slip by."""
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        assert tender.tax_percent > 0  # otherwise this test can't distinguish the two bases
        _snapshot, group = _snapshot_and_group(session, tender, "M/s Awan Tech")
        supplier = session.get(Supplier, group.supplier_id)
        assert group.store_value != group.contract_value

        content = generate_contract_award(tender, group, supplier, contract_no="C-1", rules=DEFAULT_RULES)
        full_text = _full_text(Document(BytesIO(content)))

        store_based_deposit = f"{group.store_value * 0.05:,.2f}"
        store_based_stamp_duty = f"{group.store_value * 0.0025:,.2f}"
        contract_based_deposit = f"{group.contract_value * 0.05:,.2f}"
        contract_based_stamp_duty = f"{group.contract_value * 0.0025:,.2f}"

        assert store_based_deposit in full_text
        assert store_based_stamp_duty in full_text
        assert contract_based_deposit not in full_text
        assert contract_based_stamp_duty not in full_text


def test_ampersand_in_firm_name_survives_rendering():
    """Regression test: docxtpl's XML patching does an 'unescape html
    entities' pass, so un-escaped '&' in context values (e.g. a firm name
    like "M/s X & Sons") corrupts not just that value but nearby XML too.
    See docx_export._esc()."""
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        snapshot = save_proposal_snapshot(session, tender.id)
        group = snapshot.firm_groups[0]
        supplier = session.get(Supplier, group.supplier_id)
        supplier.name = "M/s Test & Sons"
        group.supplier_name = "M/s Test & Sons"

        content = generate_contract_award(tender, group, supplier, contract_no="C-2", rules=DEFAULT_RULES)
        full_text = _full_text(Document(BytesIO(content)))

        assert "M/s Test & Sons" in full_text
        assert "SCHEDULE OF STORES TO BE SUPPLIED" in full_text  # static heading, unrelated to the value above
        assert "{{" not in full_text and "{%" not in full_text


# --- Purchase Proposal (PP) --------------------------------------------------


def test_purchase_proposal_doc_lists_every_firm_group():
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        proposal = build_purchase_proposal(session, tender.id)
        for group in proposal.firm_groups:
            supplier = session.get(Supplier, group.supplier_id)
            supplier.address = f"Address for {supplier.name}"
            session.add(supplier)
        session.commit()

        snapshot = save_proposal_snapshot(session, tender.id)
        suppliers_by_id = {g.supplier_id: session.get(Supplier, g.supplier_id) for g in snapshot.firm_groups}

        content = generate_purchase_proposal_doc(tender, snapshot, suppliers_by_id)
        doc = Document(BytesIO(content))
        full_text = _full_text(doc)

        assert "{{" not in full_text and "{%" not in full_text
        for group in snapshot.firm_groups:
            assert group.supplier_name in full_text
            assert f"{group.store_value:,.2f}" in full_text
            assert f"{group.contract_value:,.2f}" in full_text

        # 3 suppliers quoted on the fixture (2 won, 1 - Libra - won nothing).
        assert "3" in full_text  # participating firms count appears somewhere


def test_purchase_proposal_est_cost_uses_lpr_when_present():
    """Matches the real sample PP.doc's own convention: Est Cost is
    excl-tax (sum of lpr*qty) while Offered/Contract Value is incl-tax, so
    even a "no change" scenario (lpr == awarded rate for every item) shows
    an inc% equal to the tax rate, not 0% - verified against the sample's
    own numbers (496,531.02 / 420,789 store value = 1.18 = the 18% GST
    rate, the same relationship behind its stated "40.41% inc" figure)."""
    with _fresh_session() as session:
        tender = import_tender(CS_XLSX_PATH, session)
        proposal = build_purchase_proposal(session, tender.id)

        for group in proposal.firm_groups:
            for ai in group.items:
                ai.item.lpr = ai.awarded_rate
                session.add(ai.item)
        session.commit()

        snapshot = save_proposal_snapshot(session, tender.id)
        suppliers_by_id = {g.supplier_id: session.get(Supplier, g.supplier_id) for g in snapshot.firm_groups}

        content = generate_purchase_proposal_doc(tender, snapshot, suppliers_by_id)
        full_text = _full_text(Document(BytesIO(content)))

        assert f"{snapshot.grand_contract_value:,.2f}" in full_text
        assert f"{tender.tax_percent:.2f}% inc" in full_text
